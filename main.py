import sys
import json
import os
from datetime import datetime

# PyQt5 imports - grouped by module
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QInputDialog, QLabel,
    QMenuBar, QAction, QFileDialog, QMessageBox, QDialog, QFormLayout, QLineEdit, 
    QSpinBox, QTextEdit, QComboBox, QMenu
)
from PyQt5.QtGui import QPainter, QColor, QPen, QFont
from PyQt5.QtCore import Qt, QRectF

# Version information
try:
    from version import get_version_string
except ImportError:
    # Fallback if version.py is not available
    def get_version_string():
        return "TimePlanner v1.0.0"

# Translation system
from translation import tr, init_translation_system, change_language

# Global instance tracking
_current_app_instance = None

# Lazy imports for optional dependencies
def _import_pandas():
    try:
        import pandas as pd
        return pd
    except ImportError:
        raise ImportError("pandas is required for Excel operations")

def _import_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        raise ImportError("python-docx is required for Word document operations")


# Optimized color generation with caching
import colorsys

# Pre-computed golden angle for color harmony
GOLDEN_ANGLE = 137.508
_color_cache = {}

def get_nice_color(idx):
    """Generate harmonious colors with caching for better performance"""
    if idx in _color_cache:
        return _color_cache[idx]
    
    hue = (idx * GOLDEN_ANGLE) % 360
    rgb = colorsys.hsv_to_rgb(hue/360, 0.5, 0.95)
    color = QColor(int(rgb[0]*255), int(rgb[1]*255), int(rgb[2]*255))
    
    # Cache for future use
    _color_cache[idx] = color
    return color

class AppSettings:
    """Optimized application settings with lazy loading and better error handling"""
    __slots__ = ('_user_name', '_player_number', '_requirements', '_team', '_language', 'settings_file', '_loaded')
    
    def __init__(self):
        self.settings_file = "settings.json"
        self._loaded = False
        # Initialize with defaults
        self._user_name = None
        self._player_number = None
        self._requirements = None
        self._team = None
        self._language = None
    
    def _ensure_loaded(self):
        """Lazy loading of settings"""
        if not self._loaded:
            self._load_settings()
            self._loaded = True
    
    def _load_settings(self):
        """Load settings from file with proper error handling"""
        if not os.path.exists(self.settings_file):
            return
        
        try:
            with open(self.settings_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self._user_name = data.get('user_name')
                self._player_number = data.get('player_number')
                self._requirements = data.get('requirements')
                self._team = data.get('team')
                self._language = data.get('language')
        except (json.JSONDecodeError, FileNotFoundError, PermissionError) as e:
            print(f"Warning: Could not load settings: {e}")
    
    @property
    def user_name(self):
        self._ensure_loaded()
        return self._user_name or tr("default_trainer_name")
    
    @user_name.setter
    def user_name(self, value):
        self._ensure_loaded()
        self._user_name = value
    
    @property
    def player_number(self):
        self._ensure_loaded()
        return self._player_number or 20
    
    @player_number.setter
    def player_number(self, value):
        self._ensure_loaded()
        self._player_number = value
    
    @property
    def requirements(self):
        self._ensure_loaded()
        return self._requirements or tr("default_requirements")
    
    @requirements.setter
    def requirements(self, value):
        self._ensure_loaded()
        self._requirements = value
    
    @property
    def team(self):
        self._ensure_loaded()
        return self._team or tr("default_team")
    
    @team.setter
    def team(self, value):
        self._ensure_loaded()
        self._team = value
    
    @property
    def language(self):
        self._ensure_loaded()
        return self._language or "de-de"
    
    @language.setter
    def language(self, value):
        self._ensure_loaded()
        self._language = value
    
    def save_settings(self):
        """Save settings with proper error handling"""
        self._ensure_loaded()
        try:
            data = {
                'user_name': self._user_name,
                'player_number': self._player_number,
                'requirements': self._requirements,
                'team': self._team,
                'language': self._language
            }
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except (PermissionError, OSError) as e:
            print(f"Warning: Could not save settings: {e}")
            return False
        return True


class SettingsDialog(QDialog):
    def __init__(self, settings, parent=None):
        super().__init__(parent)
        self.settings = settings
        self.setWindowTitle(tr("settings_title"))
        self.setModal(True)
        self.setMinimumWidth(400)
        
        # Using simplified language switching approach
        
        layout = QFormLayout()
        
        # User name
        self.name_edit = QLineEdit(settings.user_name)
        layout.addRow(tr("settings_trainer_name"), self.name_edit)
        
        # Player number
        self.player_spin = QSpinBox()
        self.player_spin.setRange(1, 100)
        self.player_spin.setValue(settings.player_number)
        layout.addRow(tr("settings_player_count"), self.player_spin)
        
        # Requirements
        self.requirements_edit = QTextEdit(settings.requirements)
        self.requirements_edit.setMaximumHeight(80)
        layout.addRow(tr("settings_requirements"), self.requirements_edit)
        
        # Team info
        self.team_edit = QTextEdit(settings.team)
        self.team_edit.setMaximumHeight(80)
        layout.addRow(tr("settings_team_info"), self.team_edit)
        
        # Language selection
        self.language_combo = QComboBox()
        self.language_combo.addItem(tr("language_german"), "de-de")
        self.language_combo.addItem(tr("language_english"), "en-us")
        # Set current language
        current_index = self.language_combo.findData(settings.language)
        if current_index >= 0:
            self.language_combo.setCurrentIndex(current_index)
        layout.addRow(tr("settings_language"), self.language_combo)
        
        # Buttons
        button_layout = QHBoxLayout()
        ok_button = QPushButton(tr("button_ok"))
        cancel_button = QPushButton(tr("button_cancel"))
        
        ok_button.clicked.connect(self.accept)
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        
        layout.addRow(button_layout)
        self.setLayout(layout)
    
    def accept(self):
        # Check if language changed
        old_language = self.settings.language
        new_language = self.language_combo.currentData()
        
        # Save the values back to settings
        self.settings.user_name = self.name_edit.text().strip()
        self.settings.player_number = self.player_spin.value()
        self.settings.requirements = self.requirements_edit.toPlainText().strip()
        self.settings.team = self.team_edit.toPlainText().strip()
        self.settings.language = new_language
        self.settings.save_settings()
        
        # Change language if it was updated
        if old_language != new_language:
            change_language(new_language)
            # Update the main app UI if it exists
            global _current_app_instance
            if _current_app_instance:
                _current_app_instance.update_all_ui()
            # Show message to user that language was changed
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.information(self, tr("settings_title"), tr("language_changed_dynamic_message"))
        
        super().accept()
    
    def update_language(self, language_code):
        """Update dialog text when language changes"""
        self.setWindowTitle(tr("settings_title"))
        # Update form labels by recreating the layout
        self.refresh_ui()
    
    def refresh_ui(self):
        """Refresh the UI elements with new translations"""
        # Get current values
        current_name = self.name_edit.text()
        current_players = self.player_spin.value()
        current_requirements = self.requirements_edit.toPlainText()
        current_team = self.team_edit.toPlainText()
        current_language = self.language_combo.currentData()
        
        # Clear and rebuild layout
        old_layout = self.layout()
        if old_layout:
            while old_layout.count():
                child = old_layout.takeAt(0)
                if child and child.widget():
                    child.widget().deleteLater()
        
        # Rebuild layout with new translations
        layout = QFormLayout()
        
        # User name
        self.name_edit = QLineEdit(current_name)
        layout.addRow(tr("settings_trainer_name"), self.name_edit)
        
        # Player number
        self.player_spin = QSpinBox()
        self.player_spin.setRange(1, 100)
        self.player_spin.setValue(current_players)
        layout.addRow(tr("settings_player_count"), self.player_spin)
        
        # Requirements
        self.requirements_edit = QTextEdit(current_requirements)
        self.requirements_edit.setMaximumHeight(80)
        layout.addRow(tr("settings_requirements"), self.requirements_edit)
        
        # Team info
        self.team_edit = QTextEdit(current_team)
        self.team_edit.setMaximumHeight(80)
        layout.addRow(tr("settings_team_info"), self.team_edit)
        
        # Language selection
        self.language_combo = QComboBox()
        self.language_combo.addItem(tr("language_german"), "de-de")
        self.language_combo.addItem(tr("language_english"), "en-us")
        # Set current language
        current_index = self.language_combo.findData(current_language)
        if current_index >= 0:
            self.language_combo.setCurrentIndex(current_index)
        layout.addRow(tr("settings_language"), self.language_combo)
        
        # Buttons
        button_layout = QHBoxLayout()
        ok_button = QPushButton(tr("button_ok"))
        cancel_button = QPushButton(tr("button_cancel"))
        
        ok_button.clicked.connect(self.accept)
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        
        layout.addRow(button_layout)
        self.setLayout(layout)


class AddSectionDialog(QDialog):
    def __init__(self, max_duration, parent=None):
        super().__init__(parent)
        self.setWindowTitle(tr("add_section_title"))
        self.setModal(True)
        self.setMinimumWidth(400)
        
        layout = QFormLayout()
        
        # Name
        self.name_edit = QLineEdit()
        layout.addRow(tr("add_section_name"), self.name_edit)
        
        # Duration
        self.duration_spin = QSpinBox()
        self.duration_spin.setRange(1, max_duration)
        self.duration_spin.setValue(10)
        self.duration_spin.setSuffix(" min")
        layout.addRow(f"{tr('add_section_duration')} (max {max_duration} min):", self.duration_spin)
        
        # Organisation
        self.organisation_combo = QComboBox()
        self.organisation_combo.addItems([tr("organization_exercise"), tr("organization_game")])
        layout.addRow(tr("add_section_organization"), self.organisation_combo)
        
        # Explanation
        self.explanation_edit = QTextEdit()
        self.explanation_edit.setMaximumHeight(100)
        self.explanation_edit.setPlaceholderText(tr("add_section_explanation_placeholder"))
        layout.addRow(tr("add_section_explanation"), self.explanation_edit)
        
        # Tools
        self.tools_edit = QLineEdit()
        self.tools_edit.setPlaceholderText(tr("add_section_tools_placeholder"))
        layout.addRow(tr("add_section_tools"), self.tools_edit)
        
        # Buttons
        button_layout = QHBoxLayout()
        ok_button = QPushButton(tr("button_ok"))
        cancel_button = QPushButton(tr("button_cancel"))
        
        ok_button.clicked.connect(self.accept)
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        
        layout.addRow(button_layout)
        self.setLayout(layout)
    
    def get_values(self):
        return {
            'name': self.name_edit.text().strip(),
            'duration': self.duration_spin.value(),
            'organisation': self.organisation_combo.currentText(),
            'explanation': self.explanation_edit.toPlainText().strip(),
            'tools': self.tools_edit.text().strip()
        }


class EditSectionDialog(QDialog):
    def __init__(self, section, max_duration, parent=None):
        super().__init__(parent)
        self.section = section
        self.setWindowTitle(tr("edit_section_title"))
        self.setModal(True)
        self.setMinimumWidth(400)
        
        layout = QFormLayout()
        
        # Name
        self.name_edit = QLineEdit(section.name)
        layout.addRow(tr("edit_section_name"), self.name_edit)
        
        # Duration
        self.duration_spin = QSpinBox()
        self.duration_spin.setRange(1, max_duration)
        self.duration_spin.setValue(section.duration)
        self.duration_spin.setSuffix(" min")
        layout.addRow(f"{tr('edit_section_duration')} (max {max_duration} min):", self.duration_spin)
        
        # Organisation
        self.organisation_combo = QComboBox()
        self.organisation_combo.addItems([tr("organization_exercise"), tr("organization_game")])
        current_org = getattr(section, 'organisation', tr("organization_exercise"))
        # Handle both German and translated organization values
        if current_org == "Übungsform":
            current_org = tr("organization_exercise")
        elif current_org == "Spielform":
            current_org = tr("organization_game")
        self.organisation_combo.setCurrentText(current_org)
        layout.addRow(tr("edit_section_organization"), self.organisation_combo)
        
        # Explanation
        self.explanation_edit = QTextEdit()
        self.explanation_edit.setMaximumHeight(100)
        self.explanation_edit.setPlaceholderText(tr("edit_section_explanation_placeholder"))
        self.explanation_edit.setPlainText(getattr(section, 'explanation', ''))
        layout.addRow(tr("edit_section_explanation"), self.explanation_edit)
        
        # Tools
        self.tools_edit = QLineEdit()
        self.tools_edit.setPlaceholderText(tr("edit_section_tools_placeholder"))
        self.tools_edit.setText(getattr(section, 'tools', ''))
        layout.addRow(tr("edit_section_tools"), self.tools_edit)
        
        # Buttons
        button_layout = QHBoxLayout()
        ok_button = QPushButton(tr("button_ok"))
        cancel_button = QPushButton(tr("button_cancel"))
        delete_button = QPushButton(tr("button_delete"))
        delete_button.setStyleSheet("background-color: #ff6b6b; color: white;")
        
        ok_button.clicked.connect(self.accept)
        cancel_button.clicked.connect(self.reject)
        delete_button.clicked.connect(self.delete_section)
        
        button_layout.addWidget(delete_button)
        button_layout.addStretch()
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        
        layout.addRow(button_layout)
        self.setLayout(layout)
        
        self.delete_requested = False
    
    def delete_section(self):
        reply = QMessageBox.question(
            self, 
            tr("delete_section_title"), 
            tr("delete_section_confirm").format(self.section.name),
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.delete_requested = True
            self.accept()
    
    def get_values(self):
        return {
            'name': self.name_edit.text().strip(),
            'duration': self.duration_spin.value(),
            'organisation': self.organisation_combo.currentText(),
            'explanation': self.explanation_edit.toPlainText().strip(),
            'tools': self.tools_edit.text().strip()
        }


class TimeSection:
    """Optimized time section with memory-efficient slots"""
    __slots__ = ('start', 'duration', 'name', 'color', 'organisation', 'explanation', 'tools')
    
    def __init__(self, start, duration, name, color, organisation=None, explanation="", tools=""):
        self.start = start  # in minutes
        self.duration = duration  # in minutes
        self.name = name
        self.color = color
        self.organisation = organisation if organisation is not None else tr("organization_exercise")
        self.explanation = explanation  # Detailed explanation of the exercise
        self.tools = tools  # List of tools needed
    
    @property
    def end_time(self):
        """Calculate end time efficiently"""
        return self.start + self.duration
    
    def to_dict(self):
        """Convert to dictionary for serialization"""
        return {
            'start': self.start,
            'duration': self.duration,
            'name': self.name,
            'color': f'#{self.color.red():02x}{self.color.green():02x}{self.color.blue():02x}',
            'organisation': self.organisation,
            'explanation': self.explanation,
            'tools': self.tools
        }


class BarWidget(QWidget):
    """Optimized bar widget with efficient painting and memory management"""
    
    # Class constants for better performance
    MARGIN = 80
    BAR_HEIGHT = 120
    DASH_LINE_ALPHA = 120
    BACKGROUND_COLOR = QColor(250, 252, 255)
    BAR_COLOR = QColor(230, 235, 245)
    
    def __init__(self, total_minutes=120, sections=None):
        super().__init__()
        self.total_minutes = total_minutes
        self.sections = sections if sections else []
        self.setMinimumWidth(1300)
        self.setMinimumHeight(340)
        
        # Drag state variables
        self._drag_idx = None
        self._drag_offset = 0
        self._drag_x = None
        self._drag_y = None
        self._drag_section = None
        
        # Performance optimizations
        self.setMouseTracking(True)
        self.setAttribute(Qt.WA_OpaquePaintEvent, True)  # Optimize painting
        
        # Pre-create commonly used objects to reduce allocations
        self._dash_pen = QPen(QColor(120, 120, 120, self.DASH_LINE_ALPHA), 1, Qt.DashLine)
        self._text_pen = QPen(QColor(30, 30, 30))
        self._text_font = QFont()
        self._text_font.setPointSize(8)
        self._text_font.setBold(False)
        
        # Cache for paint calculations
        self._last_width = 0
        self._cached_bar_width = 0

    def paintEvent(self, event):
        """Optimized paint event with reduced object creation"""
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        # Cache calculations for performance
        width = self.width()
        height = self.height()
        
        # Update cached values if width changed
        if width != self._last_width:
            self._last_width = width
            self._cached_bar_width = width - 2 * self.MARGIN
        
        bar_y = height // 2 - self.BAR_HEIGHT // 2
        
        # Background - use cached color
        painter.setBrush(self.BACKGROUND_COLOR)
        painter.setPen(Qt.NoPen)
        painter.drawRect(0, 0, width, height)
        
        # Main bar - use cached color
        painter.setBrush(self.BAR_COLOR)
        painter.drawRoundedRect(self.MARGIN, bar_y, self._cached_bar_width, self.BAR_HEIGHT, 16, 16)
        
        # Set up text rendering once
        painter.setFont(self._text_font)
        
        # Draw segments (except dragged one)
        self._paint_sections(painter, bar_y, exclude_dragged=True)
        
        # Draw dragged segment as ghost
        self._paint_dragged_section(painter, bar_y)
    
    def _paint_sections(self, painter, bar_y, exclude_dragged=False):
        """Efficiently paint sections with minimal object creation"""
        for idx, section in enumerate(self.sections):
            if exclude_dragged and self._drag_idx == idx and self._drag_section is not None:
                continue
            
            # Calculate positions once
            x_start = self.MARGIN + (section.start / self.total_minutes) * self._cached_bar_width
            x_end = self.MARGIN + (section.end_time / self.total_minutes) * self._cached_bar_width
            
            # Draw segment
            rect = QRectF(x_start, bar_y, x_end - x_start, self.BAR_HEIGHT)
            painter.setBrush(section.color)
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(rect, 12, 12)
            
            # Draw dashed lines using pre-created pen
            painter.setPen(self._dash_pen)
            x_start_int, x_end_int = int(x_start), int(x_end)
            line_top = int(bar_y - 8)
            line_bottom = int(bar_y + self.BAR_HEIGHT + 18)
            painter.drawLine(x_start_int, line_top, x_start_int, line_bottom)
            painter.drawLine(x_end_int, line_top, x_end_int, line_bottom)
            
            # Draw text using pre-created pen and font
            painter.setPen(self._text_pen)
            text = f"{section.name} ({section.duration} min)"
            text_rect = QRectF(x_start, bar_y + self.BAR_HEIGHT + 6, x_end - x_start, 18)
            painter.drawText(text_rect, Qt.AlignCenter, text)
    
    def _paint_dragged_section(self, painter, bar_y):
        """Paint the dragged section as a ghost"""
        if not (self._drag_idx is not None and self._drag_section is not None and 
                self._drag_x is not None and self._drag_y is not None):
            return
        
        drag = self._drag_section
        drag_width = (drag.duration / self.total_minutes) * self._cached_bar_width
        rect = QRectF(self._drag_x - self._drag_offset, self._drag_y - self.BAR_HEIGHT // 2, 
                     drag_width, self.BAR_HEIGHT)
        
        # Ghost section
        painter.setBrush(drag.color.lighter(120))
        ghost_pen = QPen(QColor(120, 120, 120, 80), 2, Qt.DashLine)
        painter.setPen(ghost_pen)
        painter.drawRoundedRect(rect, 12, 12)
        
        # Ghost dashed lines
        painter.setPen(self._dash_pen)
        left_int, right_int = int(rect.left()), int(rect.right())
        top_int, bottom_int = int(rect.top() - 8), int(rect.bottom() + 18)
        painter.drawLine(left_int, top_int, left_int, bottom_int)
        painter.drawLine(right_int, top_int, right_int, bottom_int)
        
        # Ghost text
        painter.setPen(self._text_pen)
        text = f"{drag.name} ({drag.duration} min)"
        text_rect = QRectF(rect.left(), rect.bottom() + 6, rect.width(), 18)
        painter.drawText(text_rect, Qt.AlignCenter, text)

    def mousePressEvent(self, event):
        """Optimized mouse press handling"""
        if event.button() == Qt.RightButton:
            self.show_context_menu(event)
            return
        if event.button() != Qt.LeftButton:
            return
        
        # Use cached values for performance
        x, y = event.x(), event.y()
        bar_y = self.height() // 2 - self.BAR_HEIGHT // 2
        
        # Find clicked section efficiently
        clicked_section = self._find_section_at_point(x, y, bar_y)
        if clicked_section is not None:
            idx, section = clicked_section
            x_start = self.MARGIN + (section.start / self.total_minutes) * self._cached_bar_width
            
            self._drag_idx = idx
            self._drag_offset = x - x_start
            self._drag_x = x
            self._drag_y = y
            self._drag_section = section
            self.setCursor(Qt.ClosedHandCursor)
            self.update()
    
    def _find_section_at_point(self, x, y, bar_y):
        """Efficiently find section at given point"""
        # Quick bounds check
        if not (self.MARGIN <= x <= self.width() - self.MARGIN and 
                bar_y <= y <= bar_y + self.BAR_HEIGHT):
            return None
        
        for idx, section in enumerate(self.sections):
            x_start = self.MARGIN + (section.start / self.total_minutes) * self._cached_bar_width
            x_end = self.MARGIN + (section.end_time / self.total_minutes) * self._cached_bar_width
            
            if x_start <= x <= x_end:
                return idx, section
        
        return None

    def mouseMoveEvent(self, event):
        if self._drag_idx is not None and self._drag_section is not None:
            width = self.width()
            margin = 80
            bar_height = 120
            bar_width = width-2*margin
            x = event.x()
            y = event.y()
            self._drag_x = x
            self._drag_y = y
            self.setCursor(Qt.ClosedHandCursor)
            self.update()
        else:
            # Hover-Effekt
            width = self.width()
            margin = 80
            bar_height = 120
            bar_width = width-2*margin
            x = event.x()
            y = event.y()
            bar_y = self.height()//2 - bar_height//2
            hovered = None
            for idx, section in enumerate(self.sections):
                x_start = margin + (section.start/self.total_minutes)*bar_width
                x_end = margin + ((section.start+section.duration)/self.total_minutes)*bar_width
                rect = QRectF(x_start, bar_y, x_end-x_start, bar_height)
                if rect.contains(x, y):
                    hovered = idx
                    break
            if hovered is not None:
                self.setCursor(Qt.OpenHandCursor)
            else:
                self.setCursor(Qt.ArrowCursor)

    def mouseReleaseEvent(self, event):
        if self._drag_idx is not None and self._drag_section is not None:
            width = self.width()
            margin = 80
            bar_height = 120
            bar_width = width-2*margin
            x = event.x()
            # Bestimme neue Position im Array anhand Cursor-X
            rel_x = x - margin
            total = 0
            new_pos = 0
            for i, s in enumerate(self.sections):
                if i == self._drag_idx:
                    continue
                seg_len = (s.duration/self.total_minutes)*bar_width
                if rel_x < total + seg_len/2:
                    new_pos = i
                    break
                total += seg_len
            else:
                new_pos = len(self.sections)-1
            if new_pos != self._drag_idx:
                seg = self.sections.pop(self._drag_idx)
                self.sections.insert(new_pos, seg)
            # Startzeiten neu berechnen
            cur = 0
            for s in self.sections:
                s.start = cur
                cur += s.duration
            self._drag_idx = None
            self._drag_section = None
            self._drag_x = None
            self._drag_y = None
            self.setCursor(Qt.ArrowCursor)
            self.update()

    def show_context_menu(self, event):
        """Show context menu for right-click on segments"""
        width = self.width()
        margin = 80
        bar_height = 120
        bar_width = width-2*margin
        x = event.x()
        y = event.y()
        bar_y = self.height()//2 - bar_height//2
        
        # Find which section was clicked
        clicked_section = None
        for idx, section in enumerate(self.sections):
            x_start = margin + (section.start/self.total_minutes)*bar_width
            x_end = margin + ((section.start+section.duration)/self.total_minutes)*bar_width
            rect = QRectF(x_start, bar_y, x_end-x_start, bar_height)
            if rect.contains(x, y):
                clicked_section = (idx, section)
                break
        
        if clicked_section is None:
            return
        
        # Create context menu
        menu = QMenu(self)
        
        edit_action = QAction(tr("context_edit"), self)
        edit_action.triggered.connect(lambda: self.edit_section(clicked_section[0]))
        menu.addAction(edit_action)
        
        delete_action = QAction(tr("context_delete"), self)
        delete_action.triggered.connect(lambda: self.delete_section(clicked_section[0]))
        menu.addAction(delete_action)
        
        # Show menu at cursor position
        menu.exec_(event.globalPos())
    
    def edit_section(self, section_idx):
        """Edit a section"""
        if section_idx >= len(self.sections):
            return
            
        section = self.sections[section_idx]
        
        # Calculate available duration (current duration + remaining time)
        remaining_time = self.total_minutes - sum(s.duration for i, s in enumerate(self.sections) if i != section_idx)
        max_duration = remaining_time
        
        # Get parent TimePlannerApp to access the edit dialog
        parent = self.parent()
        while parent and not isinstance(parent, TimePlannerApp):
            parent = parent.parent()
        
        if parent:
            parent.edit_section_dialog(section, section_idx, max_duration)
    
    def delete_section(self, section_idx):
        """Delete a section with confirmation"""
        if section_idx >= len(self.sections):
            return
            
        section = self.sections[section_idx]
        reply = QMessageBox.question(
            self, 
            tr("delete_section_title"), 
            tr("delete_section_confirm").format(section.name),
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.sections.pop(section_idx)
            # Recalculate start times
            cur = 0
            for s in self.sections:
                s.start = cur
                cur += s.duration
            self.update()


class TimePlannerApp(QWidget):
    def __init__(self, settings=None):
        super().__init__()
        # Initialize settings
        self.settings = settings if settings is not None else AppSettings()
        
                # Set as current app instance for language updates
        global _current_app_instance
        _current_app_instance = self
        
        name, ok1 = QInputDialog.getText(self, tr("project_name_title"), tr("project_name_prompt"), text=tr("default_project_name"))
        if not ok1 or not name.strip():
            name = tr("default_project_name")
        total, ok2 = QInputDialog.getInt(self, tr("total_duration_title"), tr("total_duration_prompt"), 120, 10, 1440)
        if not ok2:
            total = 120
        self.plan_name = name.strip()
        self.total_minutes = total
        self.setWindowTitle(f"{get_version_string()} - {self.plan_name}")
        self.setStyleSheet("""
            QWidget {
                background: #f4f7fa;
            }
            QLabel#titleLabel {
                font-size: 22px;
                font-weight: bold;
                color: #2a3a5a;
                margin-bottom: 10px;
            }
            QLabel#infoLabel {
                font-size: 14px;
                color: #444;
                margin-bottom: 10px;
            }
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #6a8cff, stop:1 #3a4a7a);
                color: white;
                border-radius: 8px;
                font-size: 14px;
                padding: 8px 18px;
                margin: 8px 0;
            }
            QPushButton:hover {
                background: #4a6cff;
            }
        """)
        self.bar = BarWidget(total_minutes=self.total_minutes)
        self.title = QLabel(f"{self.plan_name}")
        self.title.setObjectName("titleLabel")
        self.info = QLabel(tr('total_duration_label').format(self.total_minutes//60, self.total_minutes%60, self.total_minutes))
        self.info.setObjectName("infoLabel")
        self.add_btn = QPushButton(tr("add_section_button"))
        self.add_btn.clicked.connect(self.add_section)
        self.scale_factor = 1.0
        # Reset Button
        self.reset_btn = QPushButton(tr("reset_button"))
        self.reset_btn.clicked.connect(self.reset_plan)
        # Skalierungs-Buttons
        self.zoom_in_btn = QPushButton("+")
        self.zoom_in_btn.setToolTip(tr("zoom_in_tooltip"))
        self.zoom_in_btn.clicked.connect(self.zoom_in)
        self.zoom_out_btn = QPushButton("–")
        self.zoom_out_btn.setToolTip(tr("zoom_out_tooltip"))
        self.zoom_out_btn.clicked.connect(self.zoom_out)
        # Menüleiste
        self.menubar = QMenuBar(self)
        file_menu = self.menubar.addMenu(tr("menu_file"))
        import_excel_action = QAction(tr("menu_import_excel"), self)
        import_excel_action.setShortcut("Ctrl+I")
        import_excel_action.triggered.connect(self.import_excel)
        file_menu.addAction(import_excel_action)
        export_excel_action = QAction(tr("menu_export_excel"), self)
        export_excel_action.setShortcut("Ctrl+E")
        export_excel_action.triggered.connect(self.export_excel)
        file_menu.addAction(export_excel_action)
        export_docx_action = QAction(tr("menu_export_docx"), self)
        export_docx_action.setShortcut("Ctrl+D")
        export_docx_action.triggered.connect(self.export_docx)
        file_menu.addAction(export_docx_action)
        export_png_action = QAction(tr("menu_export_png"), self)
        export_png_action.setShortcut("Ctrl+P")
        export_png_action.triggered.connect(self.export_png)
        file_menu.addAction(export_png_action)
        file_menu.addSeparator()
        reset_action = QAction(tr("menu_reset"), self)
        reset_action.setShortcut("Ctrl+R")
        reset_action.triggered.connect(self.reset_plan)
        file_menu.addAction(reset_action)
        file_menu.addSeparator()
        scaler_action = QAction(tr("menu_window_size"), self)
        scaler_action.setShortcut("F11")
        scaler_action.triggered.connect(self.toggle_window_scaler)
        file_menu.addAction(scaler_action)
        # Skalierung
        zoom_in_action = QAction(tr("menu_zoom_in"), self)
        zoom_in_action.setShortcut("Ctrl++")
        zoom_in_action.triggered.connect(self.zoom_in)
        file_menu.addAction(zoom_in_action)
        zoom_out_action = QAction(tr("menu_zoom_out"), self)
        zoom_out_action.setShortcut("Ctrl+-")
        zoom_out_action.triggered.connect(self.zoom_out)
        file_menu.addAction(zoom_out_action)
        
        # Settings menu
        settings_menu = self.menubar.addMenu(tr("menu_settings"))
        settings_action = QAction(tr("menu_settings_action"), self)
        settings_action.setShortcut("Ctrl+,")
        settings_action.triggered.connect(self.show_settings)
        settings_menu.addAction(settings_action)
        
        # Toolbar row for buttons (nur Abschnitt hinzufügen)
        toolbar_row = QHBoxLayout()
        toolbar_row.addWidget(self.add_btn)
        toolbar_row.addStretch(1)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 20, 30, 20)
        layout.setSpacing(10)
        layout.setMenuBar(self.menubar)
        layout.addWidget(self.title, alignment=Qt.AlignCenter)
        layout.addWidget(self.info, alignment=Qt.AlignCenter)
        layout.addWidget(self.bar, alignment=Qt.AlignCenter)
        layout.addLayout(toolbar_row)
        self.setLayout(layout)
    # Entfernt: alles außerhalb von __init__ (Toolbar und Layout)
    
    def show_settings(self):
        dialog = SettingsDialog(self.settings, self)
        dialog.exec_()
    
    def edit_section_dialog(self, section, section_idx, max_duration):
        """Show edit dialog for a section"""
        dialog = EditSectionDialog(section, max_duration, self)
        if dialog.exec_() == QDialog.Accepted:
            if dialog.delete_requested:
                # Delete the section
                self.bar.sections.pop(section_idx)
                # Recalculate start times
                cur = 0
                for s in self.bar.sections:
                    s.start = cur
                    cur += s.duration
                self.bar.update()
            else:
                # Update the section
                values = dialog.get_values()
                if not values['name']:
                    QMessageBox.warning(self, tr("error_title"), tr("error_empty_section_name"))
                    return
                
                section.name = values['name']
                section.duration = values['duration']
                section.organisation = values['organisation']
                section.explanation = values['explanation']
                section.tools = values['tools']
                
                # Recalculate start times
                cur = 0
                for s in self.bar.sections:
                    s.start = cur
                    cur += s.duration
                self.bar.update()
    
    def reset_plan(self):
        from PyQt5.QtWidgets import QMessageBox
        reply = QMessageBox.question(self, tr("reset_title"), tr("reset_confirm"), QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.bar.sections.clear()
            self.bar.update()

    def import_excel(self):
        """Optimized Excel import with better error handling"""
        path, _ = QFileDialog.getOpenFileName(self, tr("import_excel_title"), "", tr("excel_files_filter"))
        if not path:
            return
        
        try:
            pd = _import_pandas()
            df = pd.read_excel(path, engine='openpyxl')  # Specify engine for better performance
            
            if df.empty:
                QMessageBox.warning(self, tr("error_title"), "Excel file is empty")
                return
            
            sections = []
            cur_start = 0
            
            # Define column mappings for better compatibility
            column_mappings = {
                'name': ['Abschnitt', 'Section', 'Name'],
                'duration': ['Dauer (min)', 'Duration (min)', 'Duration'],
                'color': ['Farbe', 'Color'],
                'organisation': ['Organisation', 'Organization'],
                'explanation': ['Erklärung', 'Explanation'],
                'tools': ['Hilfsmittel', 'Tools']
            }
            
            for _, row in df.iterrows():
                # Use helper function to get column value
                name = self._get_column_value(row, column_mappings['name'], '')
                duration = self._get_column_value(row, column_mappings['duration'], 0)
                color_hex = self._get_column_value(row, column_mappings['color'], '#cccccc')
                organisation = self._get_column_value(row, column_mappings['organisation'], tr('organization_exercise'))
                explanation = self._get_column_value(row, column_mappings['explanation'], '')
                tools = self._get_column_value(row, column_mappings['tools'], '')
                
                # Validate data
                if not name.strip():
                    continue  # Skip empty rows
                
                try:
                    duration = int(duration)
                    if duration <= 0:
                        continue
                except (ValueError, TypeError):
                    continue
                
                # Validate color
                try:
                    color = QColor(color_hex)
                    if not color.isValid():
                        color = get_nice_color(len(sections))
                except:
                    color = get_nice_color(len(sections))
                
                sections.append(TimeSection(cur_start, duration, str(name).strip(), color, 
                                          str(organisation), str(explanation), str(tools)))
                cur_start += duration
            
            if not sections:
                QMessageBox.warning(self, tr("error_title"), "No valid sections found in Excel file")
                return
            
            self.bar.sections = sections
            self.bar.update()
            QMessageBox.information(self, tr("import_success_title"), tr("import_success_message").format(path))
            
        except ImportError as e:
            QMessageBox.warning(self, tr("error_title"), f"Required library not found: {e}")
        except Exception as e:
            QMessageBox.warning(self, tr("error_title"), tr("import_error_message").format(str(e)))
    
    def _get_column_value(self, row, possible_columns, default):
        """Helper to get column value with multiple possible column names"""
        for col in possible_columns:
            if col in row and pd.notna(row[col]):
                return row[col]
        return default

    def zoom_in(self):
        self.scale_factor = min(self.scale_factor + 0.1, 2.5)
        self.apply_scale()

    def zoom_out(self):
        self.scale_factor = max(self.scale_factor - 0.1, 0.5)
        self.apply_scale()

    def apply_scale(self):
        # BarWidget skalieren
        w = int(1300 * self.scale_factor)
        h = int(340 * self.scale_factor)
        self.bar.setMinimumWidth(w)
        self.bar.setMinimumHeight(h)
        self.bar.resize(w, h)
        font = self.title.font()
        font.setPointSize(int(22 * self.scale_factor))
        self.title.setFont(font)
        font2 = self.info.font()
        font2.setPointSize(int(14 * self.scale_factor))
        self.info.setFont(font2)
        self.bar.update()

    def toggle_window_scaler(self):
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def export_png(self):
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        path, _ = QFileDialog.getSaveFileName(self, tr("export_png_title"), f"{self.plan_name}.png", tr("png_files_filter"))
        if not path:
            return
        pixmap = self.bar.grab()
        if pixmap.save(path, "PNG"):
            QMessageBox.information(self, tr("export_success_title"), tr("export_success_message").format(path))
        else:
            QMessageBox.warning(self, tr("error_title"), tr("export_error_message"))

    def export_excel(self):
        """Optimized Excel export with better performance and error handling"""
        path, _ = QFileDialog.getSaveFileName(self, tr("export_excel_title"), f"{self.plan_name}.xlsx", tr("excel_files_filter"))
        if not path:
            return
        
        if not self.bar.sections:
            QMessageBox.warning(self, tr("error_title"), "No sections to export")
            return
        
        try:
            pd = _import_pandas()
            
            # Pre-allocate data list for better performance
            data = []
            # Note: Python lists don't have reserve, but we can still optimize
            
            for s in self.bar.sections:
                # Use more efficient color formatting
                color_hex = f'#{s.color.red():02x}{s.color.green():02x}{s.color.blue():02x}'
                
                data.append({
                    tr('excel_column_section'): s.name,
                    tr('excel_column_start'): s.start,
                    tr('excel_column_duration'): s.duration,
                    tr('excel_column_color'): color_hex,
                    tr('excel_column_organization'): s.organisation,
                    tr('excel_column_explanation'): s.explanation,
                    tr('excel_column_tools'): s.tools
                })
            
            df = pd.DataFrame(data)
            
            # Use efficient writer with optimizations
            with pd.ExcelWriter(path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Sections')
            
            QMessageBox.information(self, tr("export_success_title"), tr("export_success_message").format(path))
            
        except ImportError as e:
            QMessageBox.warning(self, tr("error_title"), f"Required library not found: {e}")
        except PermissionError:
            QMessageBox.warning(self, tr("error_title"), "Cannot write to file. Please check if the file is open in another application.")
        except Exception as e:
            QMessageBox.warning(self, tr("error_title"), tr("export_error_message").format(str(e)))

    def export_docx(self):
        """Optimized DOCX export with lazy loading and better error handling"""
        # Check if sections exist
        if not self.bar.sections:
            QMessageBox.warning(self, tr("error_title"), "No sections to export")
            return
        
        try:
            Document = _import_docx()
        except ImportError:
            QMessageBox.warning(self, tr("error_title"), tr("docx_missing_library"))
            return
        
        # Use the specific template file
        template_path = "Trainingseinheit_Name_StandardFile_Date.docx"
        if not os.path.exists(template_path):
            QMessageBox.warning(self, tr("error_title"), tr("docx_template_not_found").format(template_path))
            return
        
        # Generate filename efficiently
        trainer_name = self._sanitize_filename(self.settings.user_name)
        theme = self._sanitize_filename(self.plan_name)
        date_str = datetime.now().strftime('%Y-%m-%d')
        default_filename = f"{tr('docx_filename_prefix')}_{trainer_name}_{theme}_{date_str}.docx"
        
        output_path, _ = QFileDialog.getSaveFileName(
            self, 
            tr("export_docx_title"), 
            default_filename, 
            tr("docx_files_filter")
        )
        if not output_path:
            return
        
        try:
            self._fill_docx_template(template_path, output_path, Document)
            QMessageBox.information(self, tr("export_success_title"), tr("export_success_message").format(output_path))
        except PermissionError:
            QMessageBox.warning(self, tr("error_title"), "Cannot write to file. Please check if the file is open in another application.")
        except Exception as e:
            QMessageBox.warning(self, tr("error_title"), tr("docx_export_error").format(str(e)))
    
    def _fill_docx_template(self, template_path, output_path, Document):
        """Optimized DOCX template filling with better performance"""
        try:
            from docx.shared import Pt
        except ImportError:
            # Fallback if docx.shared is not available
            Pt = lambda x: x
        
        doc = Document(template_path)
        
        # Generate smart tools list from all sections
        tools_combined = self._combine_tools_intelligently()
        if not tools_combined:
            tools_combined = tr('docx_no_tools')
        
        # Create replacements based on actual template placeholders
        replacements = {
            '{{Theme}}': self.plan_name,
            '{{Name}}': self.settings.user_name,
            '{{playNumber}}': str(self.settings.player_number),
            '{{tools}}': tools_combined,
            '{{Time}}': f"{self.total_minutes} {tr('docx_minutes')}",
            '{{empty}}': ""  # Clear empty placeholders
        }
        
        # Template text translations (German template content to appropriate language)
        template_text_replacements = {
            # Table headers (second table)
            'Zeit': tr('docx_table_header_time'),
            'Absicht/Ziel': tr('docx_table_header_objective'),
            'Organisation': tr('docx_table_header_organization'),
            'Übungs-/Spielform': tr('docx_table_header_exercise'),
            'Hilfsmittel': tr('docx_table_header_tools'),
            # First table labels (actual text from template)
            'Thema': tr('docx_template_thema'),
            'Vorname, Name': tr('docx_template_vorname_name'),
            'Spielerzahl:': tr('docx_template_spielerzahl'),
            'Hilfsmittel:': tr('docx_template_hilfsmittel'),
            'Mannschaft:': tr('docx_template_mannschaft'),
            'Trainingszeit in min:': tr('docx_template_trainingszeit'),
            'Voraussetzungen:': tr('docx_template_voraussetzungen'),
            'Bezirksliga-Mannschaft/U18': tr('docx_template_default_team')
        }
        
        # Replace in all table cells and apply font formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Replace placeholders first
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)
                        # Then replace template text for language support
                        for german_text, translated_text in template_text_replacements.items():
                            if german_text in paragraph.text:
                                paragraph.text = paragraph.text.replace(german_text, translated_text)
                        # Apply font formatting to all paragraphs
                        self._apply_font_formatting(paragraph)
        
        # Also apply font formatting and translations to regular paragraphs (not in tables)
        for paragraph in doc.paragraphs:
            # Apply template text translations to regular paragraphs too
            for german_text, translated_text in template_text_replacements.items():
                if german_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(german_text, translated_text)
            self._apply_font_formatting(paragraph)
        
        # Handle the sections table (Table 2)
        if len(doc.tables) >= 2:
            sections_table = doc.tables[1]  # Second table for sections
            
            # Remove the existing template row (row 1, keep header row 0)
            if len(sections_table.rows) > 1:
                # Remove template row
                sections_table._element.remove(sections_table.rows[1]._element)
            
            # Add new rows for each section
            for i, section in enumerate(self.bar.sections):
                row_cells = sections_table.add_row().cells
                if len(row_cells) >= 5:  # Ensure we have enough columns
                    # Format duration instead of time slot
                    duration_text = f"{section.duration} min"
                    
                    # Fill the cells
                    row_cells[0].text = duration_text              # Zeit (Duration)
                    row_cells[1].text = section.name               # Absicht/Ziel
                    row_cells[2].text = section.organisation       # Organisation
                    row_cells[3].text = section.explanation        # Exercise/Game form
                    row_cells[4].text = section.tools             # Hilfsmittel
                    
                    # Apply font formatting to the new row
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            self._apply_font_formatting(paragraph)
        
        doc.save(output_path)

    def _combine_tools_intelligently(self):
        """
        Optimized intelligent tool combination with better performance
        Example: "3 Balls", "6 Balls", "12 Balls" -> "12 Balls"
        """
        import re
        
        if not self.bar.sections:
            return tr('docx_no_tools')
        
        tool_quantities = {}  # tool_name -> max_quantity
        
        # Pre-compile regex for better performance
        quantity_pattern = re.compile(r'^(\d+)\s+(.+)$')
        
        for section in self.bar.sections:
            if not section.tools:
                continue
            
            # Process tools more efficiently
            tools_list = [tool.strip() for tool in section.tools.split(',') if tool.strip()]
            
            for tool in tools_list:
                match = quantity_pattern.match(tool)
                
                if match:
                    quantity = int(match.group(1))
                    tool_name = match.group(2).strip()
                    
                    # Use dict.get for cleaner code
                    tool_quantities[tool_name] = max(tool_quantities.get(tool_name, 0), quantity)
                else:
                    # No quantity found, treat as single item
                    tool_name = tool
                    if tool_name not in tool_quantities:
                        tool_quantities[tool_name] = 0  # 0 means no quantity shown
        
        if not tool_quantities:
            return tr('docx_no_tools')
        
        # Build final tools list efficiently
        final_tools = []
        for tool_name in sorted(tool_quantities.keys()):
            quantity = tool_quantities[tool_name]
            if quantity > 0:
                final_tools.append(f"{quantity} {tool_name}")
            else:
                final_tools.append(tool_name)
        
        return ', '.join(final_tools)

    def _apply_font_formatting(self, paragraph):
        """Apply Calibri font with size 14 to all runs in a paragraph"""
        from docx.shared import Pt
        
        # If paragraph has no runs but has text, we need to create a run
        if not paragraph.runs and paragraph.text:
            # Clear the paragraph and add a new run with the text
            text_content = paragraph.text
            paragraph.clear()
            run = paragraph.add_run(text_content)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
        else:
            # Apply formatting to existing runs
            for run in paragraph.runs:
                if run.font:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(14)
                    
        # Ensure paragraph has the correct font style
        if paragraph.style and hasattr(paragraph.style, 'font'):
            paragraph.style.font.name = 'Calibri'
            paragraph.style.font.size = Pt(14)

    def _sanitize_filename(self, filename):
        """
        Optimized filename sanitization with better performance
        """
        import re
        
        if not filename:
            return 'Unknown'
        
        # Pre-compile regex patterns for better performance
        if not hasattr(self, '_filename_patterns'):
            self._filename_patterns = {
                'unsafe': re.compile(r'[<>:"/\\|?*]'),
                'non_word': re.compile(r'[^\w\-_.]'),
                'multiple_hyphens': re.compile(r'-+'),
            }
        
        patterns = self._filename_patterns
        
        # Process in one pass for better performance
        sanitized = filename.replace(' ', '-')
        sanitized = patterns['unsafe'].sub('', sanitized)
        sanitized = patterns['non_word'].sub('-', sanitized)
        sanitized = patterns['multiple_hyphens'].sub('-', sanitized)
        sanitized = sanitized.strip('-')
        
        return sanitized if sanitized else 'Unknown'

    def add_section(self):
        # Ermittle das Ende des letzten Segments
        if self.bar.sections:
            last = self.bar.sections[-1]
            start = last.start + last.duration
        else:
            start = 0
        max_duration = self.bar.total_minutes - start
        if max_duration <= 0:
            QMessageBox.warning(self, "Fehler", "Keine Zeit mehr verfügbar für weitere Abschnitte.")
            return
        
        dialog = AddSectionDialog(max_duration, self)
        if dialog.exec_() == QDialog.Accepted:
            values = dialog.get_values()
            if not values['name']:
                QMessageBox.warning(self, tr("error_title"), tr("error_empty_section_name"))
                return
            
            color = get_nice_color(len(self.bar.sections))
            section = TimeSection(
                start=start,
                duration=values['duration'],
                name=values['name'],
                color=color,
                organisation=values['organisation'],
                explanation=values['explanation'],
                tools=values['tools']
            )
            self.bar.sections.append(section)
            self.bar.update()
    
    def update_language(self, language_code):
        """Update main application text when language changes"""
        # Update window title
        self.setWindowTitle(f"{get_version_string()} - {self.plan_name}")
        
        # Update info label
        self.info.setText(tr('total_duration_label').format(self.total_minutes//60, self.total_minutes%60, self.total_minutes))
        
        # Update buttons
        self.add_btn.setText(tr("add_section_button"))
        self.reset_btn.setText(tr("reset_button"))
        
        # Update tooltips
        self.zoom_in_btn.setToolTip(tr("zoom_in_tooltip"))
        self.zoom_out_btn.setToolTip(tr("zoom_out_tooltip"))
        
        # Update menu items
        self.update_menus()
    
    def update_menus(self):
        """Update menu text"""
        # Clear and rebuild menus
        self.menubar.clear()
        
        # File menu
        file_menu = self.menubar.addMenu(tr("menu_file"))
        import_excel_action = QAction(tr("menu_import_excel"), self)
        import_excel_action.setShortcut("Ctrl+I")
        import_excel_action.triggered.connect(self.import_excel)
        file_menu.addAction(import_excel_action)
        export_excel_action = QAction(tr("menu_export_excel"), self)
        export_excel_action.setShortcut("Ctrl+E")
        export_excel_action.triggered.connect(self.export_excel)
        file_menu.addAction(export_excel_action)
        export_docx_action = QAction(tr("menu_export_docx"), self)
        export_docx_action.setShortcut("Ctrl+D")
        export_docx_action.triggered.connect(self.export_docx)
        file_menu.addAction(export_docx_action)
        export_png_action = QAction(tr("menu_export_png"), self)
        export_png_action.setShortcut("Ctrl+P")
        export_png_action.triggered.connect(self.export_png)
        file_menu.addAction(export_png_action)
        file_menu.addSeparator()
        reset_action = QAction(tr("menu_reset"), self)
        reset_action.setShortcut("Ctrl+R")
        reset_action.triggered.connect(self.reset_plan)
        file_menu.addAction(reset_action)
        file_menu.addSeparator()
        scaler_action = QAction(tr("menu_window_size"), self)
        scaler_action.setShortcut("F11")
        scaler_action.triggered.connect(self.toggle_window_scaler)
        file_menu.addAction(scaler_action)
        # Zoom actions
        zoom_in_action = QAction(tr("menu_zoom_in"), self)
        zoom_in_action.setShortcut("Ctrl++")
        zoom_in_action.triggered.connect(self.zoom_in)
        file_menu.addAction(zoom_in_action)
        zoom_out_action = QAction(tr("menu_zoom_out"), self)
        zoom_out_action.setShortcut("Ctrl+-")
        zoom_out_action.triggered.connect(self.zoom_out)
        file_menu.addAction(zoom_out_action)
        
        # Settings menu
        settings_menu = self.menubar.addMenu(tr("menu_settings"))
        settings_action = QAction(tr("menu_settings_action"), self)
        settings_action.setShortcut("Ctrl+,")
        settings_action.triggered.connect(self.show_settings)
        settings_menu.addAction(settings_action)
    
    def update_all_ui(self):
        """Update all UI elements when language changes"""
        # Update window title
        self.setWindowTitle(f"{get_version_string()} - {self.plan_name}")
        
        # Update info label
        self.info.setText(tr('total_duration_label').format(self.total_minutes//60, self.total_minutes%60, self.total_minutes))
        
        # Update buttons
        self.add_btn.setText(tr("add_section_button"))
        self.reset_btn.setText(tr("reset_button"))
        
        # Update tooltips
        self.zoom_in_btn.setToolTip(tr("zoom_in_tooltip"))
        self.zoom_out_btn.setToolTip(tr("zoom_out_tooltip"))
        
        # Update menus by clearing and rebuilding them
        self.update_menus()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Load settings to get language preference
    settings = AppSettings()
    
    # Initialize translation system with user's preferred language
    init_translation_system(settings.language)
    
    # Language manager is not needed with simplified approach
    
    window = TimePlannerApp(settings)
    window.resize(1500, 500)
    window.show()
    sys.exit(app.exec_())
