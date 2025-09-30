from docx import Document
from docx.shared import Inches

def create_sample_template():
    # Create a new document
    doc = Document()

    # Add title
    title = doc.add_heading('Training Plan Template', 0)

    # Add general information section
    doc.add_heading('General Information', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Theme: ').bold = True
    p1.add_run('{{theme}}')

    p2 = doc.add_paragraph()
    p2.add_run('Trainer: ').bold = True
    p2.add_run('{{name}}')

    p3 = doc.add_paragraph()
    p3.add_run('Number of Players: ').bold = True
    p3.add_run('{{playNumber}}')

    p4 = doc.add_paragraph()
    p4.add_run('Total Duration: ').bold = True
    p4.add_run('{{time}}')

    p5 = doc.add_paragraph()
    p5.add_run('Requirements: ').bold = True
    p5.add_run('{{requirements}}')

    p6 = doc.add_paragraph()
    p6.add_run('Team: ').bold = True
    p6.add_run('{{team}}')

    p7 = doc.add_paragraph()
    p7.add_run('Tools needed: ').bold = True
    p7.add_run('{{tools}}')

    # Add sections table heading
    doc.add_heading('Training Sections', level=1)

    # Create table for sections
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time Slot'
    hdr_cells[1].text = 'Goal'
    hdr_cells[2].text = 'Organisation'
    hdr_cells[3].text = 'Explanation'
    hdr_cells[4].text = 'Tools'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Save the template
    doc.save('sample_template.docx')
    print('Sample template created: sample_template.docx')

if __name__ == "__main__":
    create_sample_template()