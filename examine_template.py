from docx import Document
import sys

try:
    doc = Document('Trainingseinheit_Name_StandardFile_Date.docx')
    print('Template loaded successfully!')
    print(f'Number of tables: {len(doc.tables)}')
    
    for i, table in enumerate(doc.tables):
        print(f'Table {i+1}:')
        print(f'  Rows: {len(table.rows)}')
        print(f'  Columns: {len(table.columns)}')
        
        if len(table.rows) > 0:
            print('  First row content:')
            for j, cell in enumerate(table.rows[0].cells):
                print(f'    Col {j+1}: "{cell.text.strip()}"')
    
    print('All paragraphs containing placeholders:')
    for i, para in enumerate(doc.paragraphs):
        if '{{' in para.text and '}}' in para.text:
            print(f'  Paragraph {i+1}: "{para.text}"')
            
    # Also check table cells for placeholders
    print('Table cells containing placeholders:')
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if '{{' in cell.text and '}}' in cell.text:
                    print(f'  Table {table_idx+1}, Row {row_idx+1}, Col {cell_idx+1}: "{cell.text}"')
                    
except Exception as e:
    print(f'Error: {e}')