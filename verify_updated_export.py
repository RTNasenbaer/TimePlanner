from docx import Document

try:
    doc = Document('test_updated_export.docx')
    print('=== UPDATED DOCX EXPORT VERIFICATION ===')
    print()
    
    print('TABLE 1 (General Information) - Tools Row:')
    if len(doc.tables) >= 1:
        table1 = doc.tables[0]
        # Look for the tools row (should be row 3 or 4)
        for i, row in enumerate(table1.rows):
            for j, cell in enumerate(row.cells):
                if 'Bälle' in cell.text or 'Hütchen' in cell.text:
                    print(f'  Row {i+1}, Col {j+1}: "{cell.text.strip()}"')
                    break
    
    print()
    print('TABLE 2 (Sections) - Duration Display:')
    if len(doc.tables) >= 2:
        table2 = doc.tables[1]
        for i, row in enumerate(table2.rows):
            if i == 0:
                print(f'  Header: Zeit = "{row.cells[0].text.strip()}"')
            else:
                duration = row.cells[0].text.strip()
                name = row.cells[1].text.strip()
                print(f'  {name}: {duration}')
                
except Exception as e:
    print(f'Error: {e}')
    import traceback
    traceback.print_exc()